from flask import Flask, request, jsonify
from flask_cors import CORS
from PyPDF2 import PdfReader
from docx import Document
from dotenv import load_dotenv
import os
import json
import io

load_dotenv()

app = Flask(__name__)
CORS(app)

client_groq = None

def get_groq_client():
    global client_groq
    if client_groq is None:
        from groq import Groq
        client_groq = Groq(api_key=os.getenv("GROQ_API_KEY"))
    return client_groq

def extract_text_from_file(file):
    filename = file.filename.lower()
    
    if filename.endswith('.pdf'):
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    
    elif filename.endswith('.docx'):
        file_bytes = io.BytesIO(file.read())
        doc = Document(file_bytes)
        text = ""
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text += para.text + "\n"
        
        return text
    
    else:
        return ""

def analyze_resume(resume_text, job_description):
    client = get_groq_client()
    prompt = f"""
You are an expert HR analyst and career coach. Analyze the resume against the job description and return ONLY a JSON object with no extra text.

Resume:
{resume_text}

Job Description:
{job_description}

Return this exact JSON structure:
{{
    "match_score": <number 0-100>,
    "matched_skills": [<list of skills from resume that match job>],
    "missing_skills": [<list of skills required but missing from resume>],
    "strengths": [<3-4 strong points of the candidate>],
    "improvements": [<3-4 specific suggestions to improve the resume>],
    "summary": "<2-3 sentence overall assessment>"
}}
"""
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
    )
    
    raw = response.choices[0].message.content.strip()
    raw = raw.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)

@app.route("/analyze", methods=["POST"])
def analyze():
    try:
        job_description = request.form.get("job_description", "")
        
        if not job_description:
            return jsonify({"error": "Job description is required"}), 400
        
        if "resume" not in request.files:
            return jsonify({"error": "Resume file is required"}), 400
        
        file = request.files["resume"]
        
        if file.filename == "":
            return jsonify({"error": "No file selected"}), 400
        
        if not (file.filename.lower().endswith('.pdf') or file.filename.lower().endswith('.docx')):
            return jsonify({"error": "Only PDF and DOCX files are supported"}), 400
        
        resume_text = extract_text_from_file(file)
        
        if not resume_text.strip():
            return jsonify({"error": "Could not extract text from file"}), 400
        
        result = analyze_resume(resume_text, job_description)
        return jsonify(result)
    
    except json.JSONDecodeError:
        return jsonify({"error": "AI response parsing failed, try again"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "running"})

if __name__ == "__main__":
    app.run(debug=True, port=5000)